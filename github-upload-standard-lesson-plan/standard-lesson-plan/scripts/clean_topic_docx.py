#!/usr/bin/env python3
import argparse
import re
from pathlib import Path

from docx import Document


def remove_element(element):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def sanitize_filename(name):
    name = re.sub(r"[\\/:*?\"<>|]", "", name).strip()
    return name or "教案"


def infer_topic(doc):
    if not doc.tables:
        return "教案"
    table = doc.tables[-1]
    row_texts = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
    for text in row_texts:
        if "《" in text and "》" in text:
            match = re.search(r"《([^》]+)》", text)
            if match:
                return match.group(1)
    candidates = [
        text for text in row_texts
        if text and text not in {"课  题", "课题", "授课时间"} and "年" not in text
    ]
    text = candidates[0] if candidates else (table.cell(0, 1).text.strip() if table.rows else "")
    match = re.search(r"《([^》]+)》", text)
    if match:
        return match.group(1)
    text = re.sub(r"第\s*\d+\s*课", "", text)
    text = re.sub(r"[（(].*?[）)]", "", text)
    text = text.replace("课题", "").replace("：", "").strip()
    return text or "教案"


def clean_docx(src, out_dir=None, output=None):
    src = Path(src)
    doc = Document(src)
    body = doc._body._element

    if len(doc.tables) >= 2:
        remove_element(doc.tables[0]._element)

    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if text in {"两江新区实验中学“三段五学六有”思维型生命课堂", "备  课  指  南"}:
            remove_element(paragraph._element)
        elif "AI赋能说明" in text:
            remove_element(paragraph._element)
        elif text == "" and doc.tables:
            try:
                if body.index(paragraph._element) < body.index(doc.tables[0]._element):
                    remove_element(paragraph._element)
            except ValueError:
                pass

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in list(cell.paragraphs):
                    if "AI赋能说明" in paragraph.text:
                        remove_element(paragraph._element)

    topic = sanitize_filename(infer_topic(doc))
    if output:
        out = Path(output)
    else:
        out_base = Path(out_dir) if out_dir else src.parent
        out_base.mkdir(parents=True, exist_ok=True)
        out = out_base / f"{topic}.docx"
    doc.save(out)
    return out


def main():
    parser = argparse.ArgumentParser(description="Clean guide matrix/AI note and save as topic-named DOCX.")
    parser.add_argument("input")
    parser.add_argument("--out-dir")
    parser.add_argument("--output")
    args = parser.parse_args()
    print(clean_docx(args.input, args.out_dir, args.output))


if __name__ == "__main__":
    main()
